import unicodedata
from docx import Document

class WordAnalyzer:
    def __init__(self, filepath):
        self.document = Document(filepath)
        self.paragraphs = self.document.paragraphs


    def get_all_information(self):
        """
        includes:
        1) fonts used
        2) indent type used
        3) font color
        4) font width percentage

        """
        pass

    def get_html_for_indentation(self):
        pass

    def get_html_for_fontname(self):
        html = ""

        paragraphs = self.paragraphs
        for paragraph in paragraphs:
            if paragraph.text != "":
                previous_font = None
                is_starting_run = True
                for run in paragraph.runs:
                    current_font = run.font.name
                    if current_font is None:
                        current_font = "no_font"

                    if current_font == previous_font:
                        html += run.text
                    else:
                        if not is_starting_run:
                            html += f"</mark><mark class=\"{current_font}\">{run.text}"
                        else:
                            html += f"<mark class=\"{current_font}\">{run.text}"
                            is_starting_run = False

                        previous_font = current_font

                html += "</mark><br>"
            else:
                html += "<br>"

        html += "</mark>"
        return html

    def get_html_for_fontwidth(self):
        html = ""

        paragraphs = self.paragraphs
        for paragraph in paragraphs:
            if paragraph.text != "":
                previous_char_class = None
                is_starting_run = True
                for run in paragraph.runs:
                    for char in run.text:

                        current_char_class = unicodedata.east_asian_width(char)

                        if current_char_class == previous_char_class:
                            html += char
                        else:
                            if not is_starting_run:
                                html += f"</mark><mark class=\"{current_char_class}\">{char}"
                            else:
                                html += f"<mark class=\"{current_char_class}\">{char}"
                                is_starting_run = False

                            previous_char_class = current_char_class

                html += "</mark><br>"
            else:
                html += "<br>"

        html += "</mark>"
        return html

    def get_css_for_indentation(self):
        pass

    def get_css_for_fontname(self):
        pass

    def get_css_for_fontwidth(self):
        pass

    def map_docx_to_css(self):
        pass