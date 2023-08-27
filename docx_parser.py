from docx import Document
import unicodedata
from flask import Markup
from random import randint
# from CSS_generator import default_css_settings
from Span import DataSpan, DataParagraph, DataDocx
from create_colors import create_middle_and_pale_color, generate_random_color

default_css_settings = {"starting_characters": {" ": "space", "\t": "tab"},
                        "font_color": "black",
                        "font_size": 20,
                        "font_family": "Arial"
                        }


# three objects to fetch: 1) filters, 2) colors, 3) main data


"""make sure that this object only returns json-like data (dont't return html strings)"""

class DocxParser:
    def __init__(self, document: Document):
        self.paragraphs = document.paragraphs
        self.docx = DataDocx()
        self.all_properties = None
        self.colors = None

    def retrieve_all(self):
        """the interface to this class"""
        self._get_main_text_data_and_filters()
        self.get_colors()

        return self.docx.get_as_list(), self.all_properties.get_as_dict(), self.colors

    @staticmethod
    def _get_font_properties(font) -> dict:
        properties = {}

        # process the property one by one
        # font-name
        if name := font.name:
            properties["font-name"] = name.replace(" ", "")

        if size := font.size:
            properties["font-size"] = size.pt

        if font.underline:
            properties["font-underline"] = True

        if font.bold:
            properties["font-bold"] = True

        if color := font.color.rgb:
            properties["color"] = str(color)

        if highlight_color := font.highlight_color:
            properties["highlight_color"] = str(highlight_color.rgb)

        if font.italic:
            properties["font-italic"] = True

        return properties

    @staticmethod
    def _get_paragraph_properties(paragraph) -> dict:
        properties = {}

        if first_line_indent := paragraph.paragraph_format.first_line_indent:
            properties["first-line-indent"] = first_line_indent.pt

        if left_indent := paragraph.paragraph_format.left_indent:
            properties["left_-indent"] = left_indent.pt

        # checking starting characters
        possible_starting_characters = default_css_settings["starting_characters"]
        no_more_starting_characters = False
        starting_characters = ""
        for run in paragraph.runs:
            for char in run.text:
                if char in possible_starting_characters:
                    char_name = possible_starting_characters[char]
                    starting_characters += char_name
                else:
                    no_more_starting_characters = True
                    break

            if no_more_starting_characters:
                break

        if starting_characters:
            properties["starting-characters"] = starting_characters

        return properties

    def _get_main_text_data_and_filters(self):
        """


        """
        self.all_properties = DataSpan.all_properties
        for paragraph in self.paragraphs:
            data_paragraph = DataParagraph()

            # check for paragraph indentation
            if paragraph.text == "":
                data_paragraph.add_break()
                continue

            paragraph_props = DocxParser._get_paragraph_properties(paragraph)
            data_paragraph.add_properties(paragraph_props)

            # for inserting html non-breakble spaces later on
            if "starting-characters" in paragraph_props:
                length_of_starting_characters = len(paragraph_props["starting-characters"])
            else:
                length_of_starting_characters = 0

            # creating <span> tag for each class found in texts within this paragraph
            for run in paragraph.runs:
                font = run.font
                current_character_props = DocxParser._get_font_properties(font)

                for i, char in enumerate(run.text):
                    char_type = unicodedata.east_asian_width(char)
                    current_character_props["char-width"] = char_type
                    char = str(Markup.escape(char))

                    # html doesn't recognize the normal space (if there are more than 1) and tab characters
                    # so we need to use non-breakable space
                    if i < length_of_starting_characters:
                        if char == " ":
                            char = "&nbsp;"
                        elif char == "\t":
                            char = "&nbsp;&nbsp;&nbsp;&nbsp;"

                    if i == 0 or current_character_props != data_paragraph[-1].get_properties():
                        data_span = DataSpan()
                        data_span.add_properties(current_character_props)
                        data_span.add_text(char)
                        data_paragraph.add_span(data_span)

                    else:
                        # current_character_props == previous_character_props:
                        # we are using the same data_span object from the previous iteration
                        data_paragraph[-1].add_text(char)

            self.docx.add_paragraph(data_paragraph)

    def get_colors(self) -> None:

        # determine the number of different colors needed
        length = 0
        all_properties = self.all_properties.get_as_dict()
        for key in all_properties:
            if not isinstance(all_properties[key], bool):
                current_filter_length = len(all_properties[key])
                if current_filter_length > length:
                    length = current_filter_length


        POSSIBLE_COLORS = ["#bcebe9", "#fea895", "#f9bac6", "#ebf7bc",
                           "#e7cee9", "#fcf1b7", "#f7d7e2", "#fecfbf",
                           "#c4e2d1", "#dde6f8"]
        if length < 11:
            chosen_colors = POSSIBLE_COLORS[:length]
        else:
            chosen_colors = POSSIBLE_COLORS[:]
            for _ in range(length - 10):
                chosen_colors.append(generate_random_color())

        for idx, color in enumerate(chosen_colors):
            middle_color, pale_color = create_middle_and_pale_color(color)
            chosen_colors[idx] = [color, middle_color, pale_color]

        self.colors = chosen_colors

