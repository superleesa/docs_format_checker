from flask import Flask, render_template, request, redirect, url_for, Markup, session, make_response
from format import *  # for pdf
from werkzeug.utils import secure_filename
from docx import Document
from docx_parser import DocxParser
from json import dumps

app = Flask(__name__)
fake_database = []
app.secret_key = "test_secret_key"
app.config["LOAD_FOLDER"] = "./loaded_docx_files"
ALLOWED_EXTENSIONS = {"docx"}


@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "GET":
        return render_template("index.html")
    else:
        file = request.files["file"]
        if file and file.filename.endswith(".docx"):
            filename = secure_filename(file.filename)
            file.save(app.config["LOAD_FOLDER"] + "/" + filename)
            return redirect(url_for("analyze2", filename=filename))


# just for test (analyzes received text)
@app.route("/analyze", methods=["GET", "POST"])
def analyze():
    if request.method == "GET":
        data = fake_database[-1]
        print(data)
        res = highlight_full_size(data)
        return render_template("analyze.html", data=data, res=Markup(res))
    else:
        return redirect("/")


# @app.route("/analyze2/<filename>", methods=["GET"])
# def analyze2(filename):
#     file_path = f"{app.config['LOAD_FOLDER']}/{filename}"
#     document = Document(file_path)
#     docx_analyzer = DocxAnalyzer(document)
#     session["document_filename"] = filename
#
#     content = docx_analyzer.create_html()
#     style = docx_analyzer.get_default_style_for_the_content()
#
#     filters_and_class_names = docx_analyzer.get_html_classes_for_each_filter()
#     colors_for_annotation = docx_analyzer.get_colors_for_annotation()
#
#     print("=======INFO=====")
#     print(docx_analyzer.html_classes)
#     print("styles goes here:", style)
#     print(filters_and_class_names)
#     print(colors_for_annotation)
#     print("========ENDINFO==========")
#
#     return render_template("analyze2.html",
#                            content=Markup(content),
#                            style=Markup(style),
#                            filters_and_class_names=Markup(filters_and_class_names),
#                            colors_for_annotation=Markup(colors_for_annotation),
#                            valid_filter_displayed_names=valid_filter_displayed_names,
#                            filename=filename)

@app.route("/anayze2/get_all_information/<filename>")
def get_all_information(filename):
    filename = session["filename"]

    # load the file and create the docxanalyzer object again
    file_path = f"{app.config['LOAD_FOLDER']}/{filename}"
    document = Document(file_path)
    docx_analyzer = DocxParser(document)

    # TODO need to send this information using JSON
    return docx_analyzer.get_all_information()


@app.route("/dump")
def dump():
    filename = "random_docx3.docx"
    file_path = f"{app.config['LOAD_FOLDER']}/{filename}"
    document = Document(file_path)
    da = DocxParser(document)
    main_text, filters, colors = da.retrieve_all()

    output = [main_text, filters, colors]
    return make_response(dumps(output), 200)

@app.route("/test")
def test():
    return "HELLLLLLLLLLLLo"

